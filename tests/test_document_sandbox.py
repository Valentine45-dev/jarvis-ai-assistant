"""Fix C — doc-generator subprocess sandbox (Job Object limiter) lifetime.

The Windows Job Object limiter created the job with KILL_ON_JOB_CLOSE but kept
the handle only as a local, so it was garbage-collected on function return — the
job's last handle closed and KILL_ON_JOB_CLOSE terminated the just-assigned child
at startup. Every doc generation died as a SILENT exit-0 with no output and no
file. The fix stashes the handle on `proc` so it outlives the function.

The end-to-end test is guarded to win32 + pywin32, because that's the exact
environment where the bug bites (no pywin32 → the limiter was a harmless no-op,
which is why it "worked on the old PC"). It would FAIL before Fix C.
"""

from __future__ import annotations

import sys

import pytest

import core.handlers.document_handler as dh
from config.settings import config


def test_limiter_config_flag_defaults_on():
    # The kill switch exists and is default-ON (limiter active unless disabled).
    assert getattr(config, "doc_generator_limit_enabled", None) is True


def _win32job_available() -> bool:
    if sys.platform != "win32":
        return False
    try:
        import win32job  # noqa: F401
        return True
    except Exception:
        return False


@pytest.mark.skipif(
    not _win32job_available(),
    reason="needs win32 + pywin32 — the env where the KILL_ON_JOB_CLOSE bug bites",
)
def test_generator_produces_file_with_limiter_enabled(tmp_path):
    """With the Job Object limiter ON, a real docx script must produce the file.

    Before Fix C this returned the silent 'exit 0, no output, no file' failure
    because the limiter killed the child at startup.
    """
    pytest.importorskip("docx")
    assert config.doc_generator_limit_enabled is True  # limiter active

    target = tmp_path / "fixc.docx"
    code = (
        "from docx import Document\n"
        f"OUTPUT_PATH = r'{target}'\n"
        "d = Document()\n"
        "d.add_paragraph('fix C')\n"
        "d.save(OUTPUT_PATH)\n"
        "print('OK: ' + OUTPUT_PATH)\n"
    )
    result = dh._run_generator(code, target)

    assert result["success"], result
    assert target.exists() and target.stat().st_size > 0
    # And it's a real, openable .docx.
    from docx import Document
    assert len(Document(str(target)).paragraphs) >= 1
